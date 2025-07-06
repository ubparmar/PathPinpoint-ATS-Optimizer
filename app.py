import os
import io
import re
import PyPDF2 as pdf
import streamlit as st
from dotenv import load_dotenv
import google.generativeai as genai
from fpdf import FPDF
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import pandas as pd
from datetime import datetime
from difflib import SequenceMatcher
from PIL import Image

# 1) Page setup
st.set_page_config(page_title="PathPinpoint", page_icon="📍", layout="wide")
load_dotenv()
API_KEY = os.getenv("GOOGLE_API_KEY", "")
if not API_KEY:
    st.error("API key missing in `.env`")
genai.configure(api_key=API_KEY)

# 2) Text cleaning
def latin1_clean(text: str) -> str:
    replacements = {
        "\u2022": "-", "\u2013": "-", "\u2014": "-",
        "\u2018": "'", "\u2019": "'", "\u201c": '"',
        "\u201d": '"', "\u2026": "..."
    }
    for orig, repl in replacements.items():
        text = text.replace(orig, repl)
    # drop anything else non-Latin-1
    return text.encode("latin-1", "ignore").decode("latin-1", "ignore")
# ─── 2a) Helper to strip Markdown before output ───────────────────────────────
def strip_markdown(text: str) -> str:
    # remove Markdown headers like “## ”
    text = re.sub(r"^#{1,6}\s*", "", text, flags=re.MULTILINE)
    # remove bold/italic markers
    text = text.replace("**", "").replace("__", "")
    # normalize bullet markers to a single “• ”
    text = re.sub(r"^[\-\*\+]\s*", "• ", text, flags=re.MULTILINE)
    return text
# 3) Simple PDF for quick dumps
def generate_pdf_simple(text: str) -> bytes:
    pdf_obj = FPDF()
    pdf_obj.add_page()
    pdf_obj.set_auto_page_break(True, 15)
    pdf_obj.set_font("Arial", size=12)

    cleaned = strip_markdown(text)               # ← apply here
    for line in cleaned.splitlines():
        pdf_obj.multi_cell(0, 8, latin1_clean(line))
    return pdf_obj.output(dest="S").encode("latin-1")


# 4) Structured DOCX
def generate_structured_docx(metrics, recs, tailored, cover, interview, skill_gap, roles, salary, networking):
    # strip Markdown from each block
    recs      = strip_markdown(recs)
    tailored  = strip_markdown(tailored)
    cover     = strip_markdown(cover)
    interview = strip_markdown(interview)
    skill_gap = strip_markdown(skill_gap)
    roles     = strip_markdown(roles)
    salary    = strip_markdown(salary)
    networking= strip_markdown(networking)


    doc = Document()
    # Title
    t = doc.add_heading("PathPinpoint Full Report", level=0)
    t.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Scores Table
    doc.add_heading("ATS & Similarity Scores", level=1)
    table = doc.add_table(rows=1, cols=2)
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text = "Metric", "Value"
    for label, value in metrics:
        row = table.add_row().cells
        row[0].text, row[1].text = label, value

    # Recommendations
    doc.add_heading("Recommendations", level=1)
    for line in recs.splitlines():
        p = doc.add_paragraph(line.lstrip("- ").strip())
        p.style = "List Bullet"

    # Tailored Resume
    doc.add_heading("AI-Tailored Resume", level=1)
    for line in tailored.splitlines():
        doc.add_paragraph(line)

    # Cover Letter
    doc.add_heading("Cover Letter", level=1)
    for line in cover.splitlines():
        doc.add_paragraph(line)

    # Interview Prep Questions
    doc.add_heading("Interview Prep Questions", level=1)
    for line in interview.splitlines():
        p = doc.add_paragraph(line.lstrip("- ").strip())
        p.style = "List Number"

    # Skill Gap Analysis
    doc.add_heading("Skill Gap Analysis", level=1)
    for line in skill_gap.splitlines():
        doc.add_paragraph(line)

    # Related Roles
    doc.add_heading("Related Roles", level=1)
    for line in roles.splitlines():
        p = doc.add_paragraph(line.lstrip("- ").strip())
        p.style = "List Bullet"

    # Salary Estimate
    doc.add_heading("Salary Estimate", level=1)
    doc.add_paragraph(salary)

    # Networking Tips
    doc.add_heading("Networking Tips", level=1)
    for line in networking.splitlines():
        p = doc.add_paragraph(line.lstrip("- ").strip())
        p.style = "List Bullet"

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()

# 5) Structured PDF
def generate_structured_pdf(metrics, recs, tailored, cover, interview, skill_gap, roles, salary, networking):
    # strip Markdown first
    recs       = strip_markdown(recs)
    tailored   = strip_markdown(tailored)
    cover      = strip_markdown(cover)
    interview  = strip_markdown(interview)
    skill_gap  = strip_markdown(skill_gap)
    roles      = strip_markdown(roles)
    salary     = strip_markdown(salary)
    networking = strip_markdown(networking)
    pdf_obj = FPDF()
    pdf_obj.add_page()
    pdf_obj.set_font("Arial", "B", 16)
    pdf_obj.cell(0, 10, latin1_clean("PathPinpoint Full Report"), ln=True, align="C")
    pdf_obj.ln(5)

    # effective page width
    epw = pdf_obj.w - 2 * pdf_obj.l_margin
    col_w = epw / 2

    # Scores table
    pdf_obj.set_font("Arial", "B", 14)
    pdf_obj.cell(0, 8, latin1_clean("ATS & Similarity Scores"), ln=True)
    pdf_obj.set_font("Arial", "", 12)
    pdf_obj.cell(col_w, 8, latin1_clean("Metric"), border=1)
    pdf_obj.cell(col_w, 8, latin1_clean("Value"), border=1, ln=True)
    for label, value in metrics:
        pdf_obj.cell(col_w, 8, latin1_clean(label), border=1)
        pdf_obj.cell(col_w, 8, latin1_clean(value), border=1, ln=True)
    pdf_obj.ln(5)

    def add_section(title, body, bullet=False, numbered=False):
        pdf_obj.set_font("Arial", "B", 14)
        pdf_obj.cell(0, 8, latin1_clean(title), ln=True)
        pdf_obj.set_font("Arial", "", 12)
        for idx, line in enumerate(body.splitlines()):
            prefix = ""
            if numbered:
                prefix = f"{idx+1}. "
            elif bullet:
                prefix = "- "
            pdf_obj.multi_cell(0, 6, latin1_clean(prefix + line))
        pdf_obj.ln(3)

    add_section("Recommendations", recs, bullet=True)
    add_section("AI-Tailored Resume", tailored)
    add_section("Cover Letter", cover)
    add_section("Interview Prep Questions", interview, numbered=True)
    add_section("Skill Gap Analysis", skill_gap)
    add_section("Related Roles", roles, bullet=True)
    add_section("Salary Estimate", salary)
    add_section("Networking Tips", networking, bullet=True)

    return pdf_obj.output(dest="S").encode("latin-1")

# 6) Extract section helper
def extract_section(full, header):
    lines, out, cap = full.splitlines(), [], False
    key = header.lower()
    for ln in lines:
        txt = ln.lstrip("- ").strip()
        if cap:
            if txt.lower().startswith("- "):
                break
            out.append(txt)
        elif txt.lower().startswith(key):
            cap = True
    return "\n".join(out).strip()

# 7) Session state
if "history" not in st.session_state:
    st.session_state.history = []
for k in [
    "resume_text","jd_text","analysis","recommendations","tailored",
    "cover_letter","interview_qs","skill_gap",
    "related_roles","salary_estimate","networking_tips"
]:
    if k not in st.session_state:
        st.session_state[k] = ""

# 8) Styling & navbar
st.markdown("""
<style>
footer, header {visibility:hidden;}
body {background:#f8f9fa;color:#333;font-family:'Segoe UI',sans-serif;}
.stButton>button {background:#0d6efd;color:white;border-radius:4px;}
.stButton>button:hover {background:#0b5ed7;}
</style>
""", unsafe_allow_html=True)

tabs = ["🏠 Home", "📊 Analysis", "📜 History", "ℹ️ About Me", "📋 About Project"]
if "active_tab" not in st.session_state:
    st.session_state.active_tab = tabs[0]
choice = st.radio("", tabs, index=tabs.index(st.session_state.active_tab), horizontal=True)
st.session_state.active_tab = choice

with st.sidebar:
    st.title("Settings")
    model_choice = st.selectbox(
        "LLM Model",
        [("Gemma 3.27B","gemma-3-27b-it")],
        format_func=lambda x: x[0]
    )[1]
    st.markdown("---")
    with st.expander("GitHub"):
        st.write("[Follow on GitHub](https://github.com/ubparmar)")

# 9) Home tab
if choice == tabs[0]:
    st.title("PathPinpoint ATS Optimizer")
    jd = st.text_area("Paste Job Description", height=180)
    uploaded = st.file_uploader("Upload Resume (PDF)", type="pdf")
    with st.expander("Usage Tips"):
        st.markdown("- Provide full JD  \n- Upload clear PDF  \n- Click Analyze & Tailor then switch to Analysis")
    if st.button("Analyze & Tailor"):
        if not jd or not uploaded:
            st.error("Please supply both JD and a PDF.")
        else:
            with st.spinner("Extracting text..."):
                reader = pdf.PdfReader(uploaded)
                rt = "\n".join(p.extract_text() or "" for p in reader.pages)
            st.session_state.resume_text = rt
            st.session_state.jd_text = jd

            ats_prompt = f"""
You are an ATS. Respond in bullets:
- Job Description Match With Ats score:
- Missing Keywords:
- Profile Summary:
- Personalized suggestions for skills, keywords and achievements that can enhance the provided resume:
- Application Success Rate:
- Skill Gap Percentage:
- Suggest 3 related job titles based on the following:
Resume:
{rt}
JD:
{jd}
"""
            with st.spinner("Running ATS analysis..."):
                resp = genai.GenerativeModel(model_choice).generate_content(ats_prompt)
            st.session_state.analysis = resp.text.strip()
            st.session_state.recommendations = extract_section(st.session_state.analysis, "personalized suggestions")

            tailor_prompt = f"""
You are a professional resume writer with 10+ years experience. Using the ORIGINAL resume and JD, write a fully tailored resume:
ORIGINAL:
{rt}
JD:
{jd}
Output only the resume.
"""
            with st.spinner("Generating tailored resume..."):
                tr = genai.GenerativeModel(model_choice).generate_content(tailor_prompt).text.strip()
            st.session_state.tailored = tr

            with st.spinner("Generating additional insights..."):
                st.session_state.cover_letter = genai.GenerativeModel(model_choice).generate_content(
                    f"Write a one-page cover letter for JD:\n{jd}"
                ).text.strip()
                st.session_state.interview_qs = genai.GenerativeModel(model_choice).generate_content(
                    f"List 5 likely interview questions for JD:\n{jd}"
                ).text.strip()
                st.session_state.skill_gap = genai.GenerativeModel(model_choice).generate_content(
                    f"Compare skills to JD requirements; give Skill Gap Percentage with bullet points and calculate skill gap and no results in table maybe bulletpoints:\n{jd}"
                ).text.strip()
                st.session_state.related_roles = genai.GenerativeModel(model_choice).generate_content(
                    "Suggest 3 related job titles. By the JD and Resume you generated:\n{jd}{rt}"
                ).text.strip()
                st.session_state.salary_estimate = genai.GenerativeModel(model_choice).generate_content(
                    f"Estimate salary range in USD for JD:\n{jd}"
                ).text.strip()
                st.session_state.networking_tips = genai.GenerativeModel(model_choice).generate_content(
                    f"Provide 3 networking tips for this JD:\n{jd}"
                ).text.strip()

            st.session_state.history.append((datetime.now(), jd, st.session_state.analysis))
            st.success("Analysis complete! Go to Analysis tab.")

# 10) Analysis tab
elif choice == tabs[1]:
    if not st.session_state.analysis:
        st.info("Run an analysis first.")
    else:
        def get_val(key):
            for ln in st.session_state.analysis.splitlines():
                t = ln.lstrip("- ").strip()
                if t.lower().startswith(key.lower()):
                    num = "".join(ch for ch in t.split(":",1)[1] if ch.isdigit() or ch==".")
                    try: return float(num)
                    except: pass
            return None

        m = get_val("Job Description Match")
        s = get_val("Application Success Rate")
        g = get_val("Skill Gap Percentage")
        sim = SequenceMatcher(None, st.session_state.resume_text, st.session_state.jd_text).ratio() * 100

        metrics = []
        if m is not None: metrics.append(("Job Match %", f"{m:.1f}%"))
        if s is not None: metrics.append(("Success Rate %", f"{s:.1f}%"))
        if g is not None: metrics.append(("Skill Gap %", f"{g:.1f}%"))
        metrics.append(("Text Similarity %", f"{sim:.1f}%"))

        st.subheader("📊 ATS & Similarity Scores")
        st.table(pd.DataFrame(metrics, columns=["Metric","Value"]))

        recs = st.session_state.recommendations or st.session_state.analysis
        st.subheader("📝 Recommendations")
        st.text_area("Recommendations", recs, height=200)

        # ─── Organized download buttons ─────────────────────────────────────────
        st.subheader("📥 Download Outputs")
        col_res, col_cover, col_tail, col_full = st.columns(4)

        # Original Resume (PDF)
        col_res.download_button(
            "Resume (PDF)",
            data=generate_pdf_simple(st.session_state.resume_text),
            file_name="resume.pdf",
            mime="application/pdf"
        )

        # Cover Letter (PDF)
        col_cover.download_button(
            "Cover Letter (PDF)",
            data=generate_pdf_simple(st.session_state.cover_letter),
            file_name="cover_letter.pdf",
            mime="application/pdf"
        )

        # Tailored Resume (PDF)
        col_tail.download_button(
            "Tailored Resume (PDF)",
            data=generate_pdf_simple(st.session_state.tailored),
            file_name="tailored_resume.pdf",
            mime="application/pdf"
        )

        # Full Report (PDF & DOCX)
        full_pdf = generate_structured_pdf(
            metrics, recs,
            st.session_state.tailored,
            st.session_state.cover_letter,
            st.session_state.interview_qs,
            st.session_state.skill_gap,
            st.session_state.related_roles,
            st.session_state.salary_estimate,
            st.session_state.networking_tips
        )
        full_docx = generate_structured_docx(
            metrics, recs,
            st.session_state.tailored,
            st.session_state.cover_letter,
            st.session_state.interview_qs,
            st.session_state.skill_gap,
            st.session_state.related_roles,
            st.session_state.salary_estimate,
            st.session_state.networking_tips
        )
        col_full.download_button(
            "Full Report (PDF)",
            data=full_pdf,
            file_name="full_report.pdf",
            mime="application/pdf"
        )
        col_full.download_button(
            "Full Report (DOCX)",
            data=full_docx,
            file_name="full_report.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        # ────────────────────────────────────────────────────────────────────────

        st.subheader("✍️ AI-Tailored Resume")
        st.text_area("Tailored Resume", st.session_state.tailored, height=300)

        for title, content in [
            ("Cover Letter", st.session_state.cover_letter),
            ("Interview Prep Questions", st.session_state.interview_qs),
            ("Skill Gap Analysis", st.session_state.skill_gap),
            ("Related Roles", st.session_state.related_roles),
            ("Salary Estimate", st.session_state.salary_estimate),
            ("Networking Tips", st.session_state.networking_tips)
        ]:
            st.subheader(title)
            st.text_area("", content, height=120)

# 11) History tab
elif choice == tabs[2]:
    st.subheader("History")
    if st.session_state.history:
        df = pd.DataFrame([
            {"Time": t.strftime("%Y-%m-%d %H:%M"), "JD Preview": jd[:30] + "..."}
            for t, jd, _ in st.session_state.history
        ]).tail(10)
        st.table(df)
    else:
        st.info("No history yet.")
elif choice == tabs[3]:
# ===== About Me tab content =====

# 2) Global CSS for this section
    st.markdown("""
        <style>
        /* Section headers */
        .section-header {
        font-size: 2rem;
        font-weight: 600;
        margin-top: 2rem;
        margin-bottom: 0.25rem;
        }
        .section-underline {
        height: 3px;
        width: 100%;
        background: linear-gradient(90deg, #ff7e5f, #feb47b, #86A8E7);
        margin-bottom: 1rem;
        }

        /* Subtitle text */
        .subtitle {
        color: #8f94a1;
        font-size: 1.1rem;
        margin-bottom: 1.5rem;
        }

        /* Bio text */
        .bio-text {
        font-size: 1rem;
        line-height: 1.6;
        margin-bottom: 1.5rem;
        }

        /* Skills grid */
        .skill-grid {
        display: grid;
        grid-template-columns: repeat(auto-fit, minmax(120px, 1fr));
        gap: 0.5rem;
        margin-bottom: 2rem;
        }
        .skill-tag {
        background-color: #161b22;
        color: #c9d1d9;
        padding: 0.5rem 1rem;
        border-radius: 6px;
        text-align: center;
        font-size: 0.9rem;
        }
        </style>
    """, unsafe_allow_html=True)

    # 3) Introduction
    st.markdown('<div class="section-header">👋 Hello, I’m Urjeet Parmar</div>', unsafe_allow_html=True)
    st.markdown('<div class="section-underline"></div>', unsafe_allow_html=True)
    st.markdown('<div class="subtitle">Web Developer → AI & Data Science Enthusiast</div>', unsafe_allow_html=True)

    col1, col2 = st.columns([2, 1], gap="large")
    with col1:
        st.markdown(
            '<div class="bio-text">'
            'I’m a former freelance Full-Stack Web Developer with experience of 2+ years of building scalable applications. '
            'Now I’m diving deep into Python, machine learning, and cloud-native architectures. '
            'I love turning data into insights and crafting seamless user experiences.'
            '</div>',
            unsafe_allow_html=True
        )

        st.markdown('<div class="section-header">🔧 Core Skills</div>', unsafe_allow_html=True)
        st.markdown('<div class="section-underline"></div>', unsafe_allow_html=True)

        # Skills grid
        skills = [
            "JavaScript","HTML/CSS","Php","Laravel",
            "Python", "FastAPI", "Docker", "Kubernetes",
            "Pandas", "NumPy", "TensorFlow", "PyTorch",
            "PostgreSQL", "MongoDB", "AWS", "Git", "CI/CD"
        ]
        grid_html = '<div class="skill-grid">'
        for skill in skills:
            grid_html += f'<div class="skill-tag">{skill}</div>'
        grid_html += '</div>'
        st.markdown(grid_html, unsafe_allow_html=True)

    with col2:
        img = Image.open("images/ub.jpg")
        st.image(img, width=300, caption="Urjeet Parmar")

    # 4) Experience
    st.markdown('<div class="section-header">💼 Professional Experience</div>', unsafe_allow_html=True)
    st.markdown('<div class="section-underline"></div>', unsafe_allow_html=True)
    st.markdown("""
    **Freelance Full-Stack Developer**  
    _Jan 2019 – Dec 2022_  
    - Built responsive web applications for clients in pet e-commerce, and vape retail.  
    - Implemented CI/CD pipelines, reducing deployment time.  
    - Worked with laravel to built scalable point of sales system with accurate inventory tracking and transfers over locations.  

    """, unsafe_allow_html=True)

    # 5) Portfolio
    st.markdown('<div class="section-header">🎨 Selected Projects</div>', unsafe_allow_html=True)
    st.markdown('<div class="section-underline"></div>', unsafe_allow_html=True)
    projects = [
        {
            "title": "Resume ATS Optimizer",
            "desc": "A Streamlit app that scores & tailors your resume to any job description using AI. Provides recommendations, Customized Resume, Cover letter, Salary Estimate and Networking tips.",
            "link": "https://github.com/ubparmar/PathPinpoint-ATS-Optimizer"
        },
        {
  "title": "RECIPACE-AI",
  "desc": "A Streamlit-powered app that lets you customize 30+ parameters to auto-generate richly detailed recipes, chef’s tips, grocery lists, timers, nutrition facts, and YouTube video scripts in one click.",
  "link": "https://github.com/ubparmar/RECIPACE-AI"
}

    ]
    for proj in projects:
        st.markdown(f"**{proj['title']}**  ")
        st.markdown(proj["desc"])
        st.markdown(f"[🔗 View on GitHub]({proj['link']})")
        st.markdown("---")

    # 6) Contacts
    st.markdown('<div class="section-header">📬 Get in Touch</div>', unsafe_allow_html=True)
    st.markdown('<div class="section-underline"></div>', unsafe_allow_html=True)
    st.markdown("""
    📧 **Email:** parmarurjeet274@gmail.com  
    🔗 **LinkedIn:** [urjeet-parmar](https://ca.linkedin.com/in/urjeet-parmar-573b892aa)  
    🐙 **GitHub:** [ubparmar](https://github.com/ubparmar)  
    """, unsafe_allow_html=True)

    # 7) Footer
    st.markdown("---")
    st.markdown("© 2025 Urjeet Parmar")
elif choice == "📋 About Project":
    # 📘 About PathPinpoint
    st.subheader("📘 About PathPinpoint")
    st.markdown("Your AI-powered ATS optimizer—get actionable insights and tailored docs in seconds.")

    col1, col2 = st.columns([3, 1], gap="large")
    with col1:
        st.markdown("""
- **AI ATS Scoring**: Instantly score how well your resume matches any job description.
- **Keyword Analysis**: See which crucial terms you’re missing.
- **Personalized Feedback**: Improve language, tone, and highlight achievements.
- **Tailored Outputs**: Auto-generate a custom resume and one-page cover letter.
- **Interview Prep & Skill Gap**: Get likely interview questions and upskill recommendations.

**Quickstart Guide**  
1. Switch to **🏠 Home**.  
2. **Paste** your Job Description & **Upload** your PDF resume.  
3. Click **Analyze & Tailor** & **Hang on tight untlil it processes your request⏳**.  
4. Review scores and recommendations on the **🔍 Analysis** tab.  
5. Download all deliverables (analysis, recommendations, tailored resume, cover letter, interview prep) directly from **🔍 Analysis**.

**Project Outputs & Benefits**  
- **See Your Fit Instantly**: Get an easy-to-understand score showing how well your resume matches any job.  
- **Spot Missing Keywords**: Discover the exact terms you’re missing so you can tweak and shine.  
- **Actionable Tips**: Friendly, bullet-point advice on how to boost your skills, achievements, and tone.  
- **Fresh Resume & Cover Letter**: Let AI craft a focused, polished resume and one-page cover letter—no more blank pages.  
- **Interview Prep & Learning Path**: Receive likely interview questions plus course or certification suggestions to bridge any gaps.  
- **Explore New Roles & Salaries**: Find similar job titles you might love and get a quick salary peek.  
- **Networking Nudge**: Simple tips on who to connect with and what to say to get noticed.  

        """, unsafe_allow_html=True)

    with col2:
        from PIL import Image
        logo = Image.open("images/pathpinpoint_logo.png")
        st.image(logo, use_column_width=True, caption="PathPinpoint ATS Optimizer")

